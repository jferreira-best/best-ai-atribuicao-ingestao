#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Gera um JSONL (chunks) a partir de arquivos locais ou de um container do Azure Blob.
VERSÃO COM OCR LOCAL (RapidOCR) INTEGRADO.
"""

import os
import re
import json
import argparse
import tempfile
import shutil
from pathlib import Path
from typing import Iterable, List, Tuple, Optional, Dict
import logging

# Configuração de Log básico
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')

# dotenv (opcional)
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# ====== Azure Blob (opcional) ======
BlobServiceClient = None
AzureNamedKeyCredential = None
try:
    from azure.storage.blob import BlobServiceClient as _BSC
    from azure.core.credentials import AzureNamedKeyCredential as _ANKC
    BlobServiceClient = _BSC
    AzureNamedKeyCredential = _ANKC
except Exception:
    pass

# ====== PDF (PyMuPDF4LLM + OCR) ======
pymupdf4llm = None
fitz = None
ocr_engine = None

try:
    import pymupdf4llm
    import fitz  # PyMuPDF standard
except ImportError:
    pass

try:
    # [NOVO] Importação do OCR
    from rapidocr_onnxruntime import RapidOCR
    ocr_engine = RapidOCR()
except ImportError:
    pass

# ====== DOCX ======
docx = None
try:
    import docx as _docx
    docx = _docx
except Exception:
    pass


# ==============================
# [NOVO] FUNÇÃO DE OCR LOCAL
# ==============================
def extract_text_with_ocr_local(path: Path) -> str:
    """
    Converte PDF em imagens e roda OCR usando RapidOCR.
    """
    if not fitz or not ocr_engine:
        print("[aviso] Bibliotecas de OCR (fitz/rapidocr) não instaladas.")
        return ""

    print(f"[OCR] Iniciando leitura de imagem em: {path.name}...")
    full_text = []
    
    try:
        with fitz.open(path) as doc:
            for page in doc:
                # Renderiza a página como imagem (zoom=2 para melhor qualidade)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                
                # Roda o OCR na imagem
                result, _ = ocr_engine(img_bytes)
                
                if result:
                    # RapidOCR retorna lista de tuplas, o texto é o índice 1
                    page_text = "\n".join([line[1] for line in result])
                    full_text.append(page_text)
                    
        return "\n\n".join(full_text)
    except Exception as e:
        print(f"[erro] Falha no OCR Local de {path.name}: {e}")
        return ""


# ==============================
# Utilidades de chunking/IO
# ==============================
def chunk_text(text: str, target_chars: int = 1500, overlap: int = 200) -> Iterable[str]:
    text = (text or "").strip()
    if not text:
        return []
    chunks = []
    i = 0
    N = len(text)
    if target_chars <= 0:
        return [text]
    if overlap < 0:
        overlap = 0

    while i < N:
        j = min(N, i + target_chars)

        if j < N:
            next_newline = text.find("\n", j)
            if next_newline != -1 and (next_newline - j) < 150:
                j = next_newline + 1

        chunk = text[i:j]
        chunks.append(chunk.strip())

        if j >= N:
            break

        i = j - overlap if overlap > 0 else j
        if i < 0:
            i = 0

    return [c for c in chunks if c]


def read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        try:
            return path.read_text(encoding="latin-1", errors="ignore")
        except Exception:
            return ""


def read_pdf(path: Path) -> str:
    """
    [MODIFICADO] Tenta PyMuPDF4LLM. Se falhar (scan), usa OCR.
    """
    text = ""
    
    # 1. Tentativa Markdown (PyMuPDF4LLM)
    if pymupdf4llm:
        try:
            text = pymupdf4llm.to_markdown(str(path))
        except Exception as e:
            print(f"[warn] Falha ao ler PDF {path.name} com PyMuPDF: {e}")

    # 2. Verificação de SCAN
    # Remove espaços para contar caracteres reais
    clean_len = len(re.sub(r'\s+', '', text))
    
    if clean_len < 100:
        print(f"[info] Texto insuficiente ({clean_len} chars) em {path.name}. Tentando OCR...")
        ocr_text = extract_text_with_ocr_local(path)
        # Só substitui se o OCR achou mais coisas
        if len(ocr_text) > clean_len:
            return ocr_text

    return text


def read_docx(path: Path) -> str:
    if docx is None:
        print(f"[warn] python-docx não instalado; ignorando DOCX: {path.name}")
        return ""
    try:
        d = docx.Document(str(path))
        full_text = []

        for p in d.paragraphs:
            t = (p.text or "").strip()
            if t:
                full_text.append(t)

        if d.tables:
            full_text.append("\n\n--- TABELAS DO DOCUMENTO ---\n")
            for tbl in d.tables:
                rows_md = []
                for row in tbl.rows:
                    cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
                    if any(cells):
                        rows_md.append("| " + " | ".join(cells) + " |")

                if rows_md:
                    full_text.append("\n".join(rows_md))
                    full_text.append("\n")

        return "\n".join(full_text)
    except Exception as e:
        print(f"[warn] Falha ao ler DOCX {path.name}: {e}")
        return ""


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md", ".csv", ".log"):
        return read_txt(path)
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    return read_txt(path)


# ==============================
# Heurísticas de metadados (SEU CÓDIGO ORIGINAL)
# ==============================
YEAR_RE = re.compile(r"\b(20[2-4]\d)\b")
DATE_BR_RE = re.compile(r"\b([0-3]?\d)[/.-]([01]?\d)[/.-](20[2-4]\d)\b")


def infer_conhecimento(file_name: str) -> str:
    fn = (file_name or "").strip().upper()
    if fn.startswith("AT"): return "Atribuição de Classes (AC)"
    if fn.startswith("AC"): return "Atribuição de Classes (AC)"
    if fn.startswith("AD"): return "Avaliação de Desempenho (AD)"
    if fn.startswith("CP"): return "Confirmação de Participação (CP)"
    if fn.startswith("PEI"): return "Programa Ensino Integral (PEI)"
    if "ATRIBUI" in fn: return "Atribuição de Classes (AC)"
    if "AVALIACA" in fn or "DESEMPENHO" in fn: return "Avaliação de Desempenho (AD)"
    if "CONFIRMACAO" in fn and "PARTICIPACAO" in fn: return "Confirmação de Participação (CP)"
    if "ENSINO INTEGRAL" in fn: return "Programa Ensino Integral (PEI)"
    return "Geral"


def norm_str(s: Optional[str]) -> str:
    return (s or "").strip()


def to_iso_date(d: str, m: str, y: str) -> str:
    try:
        dd = int(d); mm = int(m); yy = int(y)
        return f"{yy:04d}-{mm:02d}-{dd:02d}"
    except Exception:
        return ""


def guess_norma_tipo(text: str, fname: str) -> str:
    s = f"{fname}\n{text}".lower()
    if "portaria conjunta" in s: return "Portaria Conjunta"
    if "portaria" in s: return "Portaria"
    if "resolução" in s or "resolucao" in s: return "Resolução"
    if "comunicado" in s: return "Comunicado"
    if "informação" in s or "informacao" in s: return "Informação"
    if "decreto" in s: return "Decreto"
    if "lei complementar" in s: return "Lei Complementar"
    return ""


def guess_orgao_emissor(text: str) -> str:
    s = (text or "").upper()
    if "SUCOR" in s and "SUPED" in s: return "SUCOR/SUPED"
    if "CGRH" in s: return "CGRH"
    for k in ("DIPES", "SUCOR", "SUPED", "SEDUC"):
        if k in s: return k
    return ""


def guess_ano_letivo(text: str, fname: str) -> str:
    ys = re.findall(YEAR_RE, fname or "")
    if ys: return ys[0]
    ys = re.findall(YEAR_RE, text or "")
    if ys:
        try: return str(max(map(int, ys)))
        except Exception: return ys[0]
    return ""


def guess_fase_processo(text: str, fname: str) -> str:
    s = f"{fname}\n{text}"
    s_low = s.lower()
    if "conferência de dados" in s_low or "conferencia de dados" in s_low: return "Conferência de Dados"
    if "credenciamento" in s_low: return "Credenciamento"
    if "alocação" in s_low or "alocacao" in s_low:
        if "inicial" in s_low: return "Alocação Inicial"
        return "Alocação"
    if "realocação" in s_low or "realocacao" in s_low: return "Realocação"
    if "transferência" in s_low or "transferencia" in s_low:
        if "pei" in s_low: return "Transferência PEI"
        return "Transferência"
    if "confirmação de participação" in s_low or "confirmacao de participacao" in s_low:
        m = re.search(r"fase\s*(\d+)", s_low)
        if m: return f"Confirmação de Participação – Fase {m.group(1)}"
        return "Confirmação de Participação"
    if "classificação" in s_low or "classificacao" in s_low: return "Classificação"
    if "inscrição" in s_low or "inscricao" in s_low: return "Inscrição"
    if "avaliação de desempenho" in s_low or "avaliacao de desempenho" in s_low:
        if "final" in s_low: return "Avaliação de Desempenho Final"
        if "parcial" in s_low: return "Avaliação de Desempenho Parcial"
        return "Avaliação de Desempenho"
    return ""


def guess_programa(text: str) -> str:
    s = (text or "").lower()
    if "programa ensino integral" in s or " pei " in s or "no pei" in s: return "PEI"
    if "ensino integral" in s: return "PEI"
    if "tempo parcial" in s: return "Tempo Parcial"
    if "eja" in s: return "EJA"
    if "ensino técnico" in s or "novotec" in s: return "Ensino Técnico / Novotec"
    return ""


def guess_publico_alvo(text: str) -> str:
    s = (text or "").lower()
    targets = []
    if "docente" in s or "professor" in s: targets.append("Docentes")
    if "diretor" in s or "gestor" in s or "coordenador" in s: targets.append("Gestores")
    if "candidato" in s or "contratado" in s: targets.append("Candidatos/Contratados")
    if not targets:
        if "pei" in s: return "PEI"
        return "Geral"
    return ", ".join(sorted(list(set(targets))))


def guess_prazos(text: str) -> Tuple[str, str]:
    m = re.search(r"de\s+([0-3]?\d/[01]?\d)(?:/(\d{4}))?\s+a\s+([0-3]?\d/[01]?\d)(?:/(\d{4}))?", text, flags=re.I)
    if m:
        d1, y1, d2, y2 = m.group(1), m.group(2), m.group(3), m.group(4)
        year_hint = y1 or y2
        if not year_hint:
            around = re.findall(YEAR_RE, text)
            year_hint = around[0] if around else ""
        try:
            d1d, d1m = d1.split("/")
            d2d, d2m = d2.split("/")
            ini = to_iso_date(d1d, d1m, year_hint) if year_hint else ""
            fim = to_iso_date(d2d, d2m, year_hint) if year_hint else ""
            return ini, fim
        except Exception: pass
    ds = re.findall(DATE_BR_RE, text)
    if len(ds) >= 2:
        (d1, m1, y1), (d2, m2, y2) = ds[0], ds[1]
        return to_iso_date(d1, m1, y1), to_iso_date(d2, m2, y2)
    return "", ""


MESES = {"janeiro": "01", "fevereiro": "02", "março": "03", "abril": "04", "maio": "05", "junho": "06", "julho": "07", "agosto": "08", "setembro": "09", "outubro": "10", "novembro": "11", "dezembro": "12"}


def guess_data_publicacao(text: str) -> str:
    m = re.search(DATE_BR_RE, text)
    if m: return to_iso_date(m.group(1), m.group(2), m.group(3))
    m_ext = re.search(r"(\d{1,2})\s+de\s+([a-zç]+)\s+de\s+(20\d{2})", text, flags=re.IGNORECASE)
    if m_ext:
        dia, mes_nome, ano = m_ext.groups()
        mes_num = MESES.get(mes_nome.lower())
        if mes_num: return f"{ano}-{mes_num}-{int(dia):02d}"
    return ""


def extract_referencias_legais(text: str) -> List[str]:
    refs = set()
    for m in re.finditer(r"(Resolu\w+.*?)\b(\d{1,4})\/(20[1-4]\d)", text, flags=re.I):
        chunk = m.group(0).strip()
        if 6 <= len(chunk) <= 140: refs.add(chunk)
    for m in re.finditer(r"(Portaria.*?)\b(\d{1,4})\/(20[1-4]\d)", text, flags=re.I):
        chunk = m.group(0).strip()
        if 6 <= len(chunk) <= 140: refs.add(chunk)
    for m in re.finditer(r"(Lei Complementar.*?)\b(\d{1,4})\/(20\d{2}|19\d{2})", text, flags=re.I):
        chunk = m.group(0).strip()
        if 6 <= len(chunk) <= 140: refs.add(chunk)
    return sorted(refs)


# ==============================
# Blob helpers (Mantidos)
# ==============================
def download_blob_dir(tmp_dir: Path, account_name: str, account_key: str, container: str, prefix: Optional[str]) -> Path:
    if BlobServiceClient is None or AzureNamedKeyCredential is None:
        raise RuntimeError("azure-storage-blob não está instalado.")
    cred = AzureNamedKeyCredential(account_name, account_key)
    bsc = BlobServiceClient(account_url=f"https://{account_name}.blob.core.windows.net", credential=cred)
    cc = bsc.get_container_client(container)
    if not cc.exists(): raise RuntimeError(f"Container inexistente: {container}")
    out_dir = tmp_dir / "input"
    out_dir.mkdir(parents=True, exist_ok=True)
    for blob in cc.list_blobs(name_starts_with=prefix or ""):
        if not blob.size or not blob.name: continue
        if not any(blob.name.lower().endswith(ext) for ext in (".pdf", ".docx", ".txt", ".md", ".csv", ".log")): continue
        target = out_dir / blob.name
        target.parent.mkdir(parents=True, exist_ok=True)
        with open(target, "wb") as f:
            f.write(cc.download_blob(blob.name).readall())
    return out_dir


def upload_jsonl(account_name: str, account_key: str, container: str, blob_path: str, local_jsonl: Path):
    if BlobServiceClient is None or AzureNamedKeyCredential is None: raise RuntimeError("azure-storage-blob não está instalado.")
    cred = AzureNamedKeyCredential(account_name, account_key)
    bsc = BlobServiceClient(account_url=f"https://{account_name}.blob.core.windows.net", credential=cred)
    cc = bsc.get_container_client(container)
    if not cc.exists(): raise RuntimeError(f"Container inexistente: {container}")
    with open(local_jsonl, "rb") as data:
        cc.upload_blob(name=blob_path, data=data, overwrite=True)


# ==============================
# Main
# ==============================
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-dir")
    ap.add_argument("--output-jsonl", required=True)
    ap.add_argument("--assunto", default=os.getenv("ASSUNTO", "atribuicao"))
    ap.add_argument("--area-interesse", default=os.getenv("AREA_INTERESSE", "conhecimento"))
    # Azure Blob
    ap.add_argument("--account-name")
    ap.add_argument("--account-key")
    ap.add_argument("--container")
    ap.add_argument("--prefix")
    ap.add_argument("--upload-jsonl")
    # Chunking
    ap.add_argument("--target-chars", type=int, default=int(os.getenv("TARGET_CHARS", "1500")))
    ap.add_argument("--overlap", type=int, default=int(os.getenv("OVERLAP", "200")))

    # Compatibilidade com script de deploy (argumentos ignorados, mas aceitos)
    ap.add_argument("--docint-endpoint")
    ap.add_argument("--docint-key")

    args = ap.parse_args()

    tmp_root = Path(tempfile.mkdtemp(prefix="mkjsonl_"))
    work_dir = None

    try:
        if args.input_dir:
            work_dir = Path(args.input_dir).resolve()
            if not work_dir.exists(): raise RuntimeError(f"Diretório não existe: {work_dir}")
        else:
            if args.container and args.account_name and args.account_key:
                work_dir = download_blob_dir(tmp_root, args.account_name, args.account_key, args.container, args.prefix)
            else:
                raise RuntimeError("Informe --input-dir OU (--container + --account-name + --account-key).")

        files = [p for p in work_dir.rglob("*") if p.is_file() and p.suffix.lower() in (".pdf", ".docx", ".txt", ".md", ".csv", ".log")]
        files.sort(key=lambda x: x.name.lower())

        records: List[dict] = []
        for f in files:
            print(f"[info] Processando: {f.name}")
            text = extract_text_from_file(f)
            
            if not text.strip():
                print(f"[ignorado] Vazio: {f.name}")
                continue

            doc_title = f.stem
            
            # [NOVO] Heurística de Correção de Título (Se OCR achou algo como "RESOLUÇÃO SEDUC")
            lines = text.split('\n')
            # Olha apenas nas primeiras 15 linhas para achar o título
            for l in lines[:15]:
                clean_l = l.strip().upper()
                if ("RESOLUÇÃO" in clean_l or "PORTARIA" in clean_l or "DECRETO" in clean_l) and len(clean_l) < 150:
                    doc_title = l.strip()
                    break

            is_gloss = bool(re.search(r"gloss[aá]rio", f.name, flags=re.I))

            # metadados
            conhecimento = infer_conhecimento(f.name)
            norma_tipo = guess_norma_tipo(text, f.name)
            orgao_emissor = guess_orgao_emissor(text)
            data_publicacao = guess_data_publicacao(text)
            ano_letivo = guess_ano_letivo(text, f.name)
            fase_processo = guess_fase_processo(text, f.name)
            programa = guess_programa(text)
            publico_alvo = guess_publico_alvo(text)
            prazo_inicio, prazo_fim = guess_prazos(text)
            referencias_legais = extract_referencias_legais(text)

            try: source_rel = f.relative_to(work_dir).as_posix()
            except: source_rel = f.name

            for idx, chunk in enumerate(chunk_text(text, target_chars=args.target_chars, overlap=args.overlap), start=1):
                rec = {
                    "id": f"{source_rel}#chunk{idx}",
                    "id_original": f"{source_rel}#chunk{idx}",
                    "source": source_rel,
                    "source_file": source_rel,
                    "doc_title": doc_title,
                    "assunto": args.assunto,
                    "area_interesse": args.area_interesse,
                    "conhecimento": conhecimento,
                    "content": chunk,
                    "text": chunk,
                    "is_glossario": is_gloss,
                    "chunk": idx,
                    "norma_tipo": norma_tipo,
                    "orgao_emissor": orgao_emissor,
                    "data_publicacao": data_publicacao,
                    "ano_letivo": ano_letivo,
                    "fase_processo": fase_processo,
                    "programa": programa,
                    "publico_alvo": publico_alvo,
                    "prazo_inicio": prazo_inicio,
                    "prazo_fim": prazo_fim,
                    "referencias_legais": referencias_legais,
                }
                records.append(rec)

        out = Path(args.output_jsonl).resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
        with out.open("w", encoding="utf-8") as w:
            for rec in records:
                w.write(json.dumps(rec, ensure_ascii=False) + "\n")

        print(f"[ok] JSONL gerado: {out} (chunks: {len(records)})")

        if args.upload_jsonl:
            upload_jsonl(args.account_name, args.account_key, args.container, args.upload_jsonl, out)
            print(f"[ok] JSONL enviado ao blob: {args.container}/{args.upload_jsonl}")

    finally:
        if (not args.input_dir) and tmp_root.exists():
            shutil.rmtree(tmp_root, ignore_errors=True)

    print("[done] Processo concluído.")

if __name__ == "__main__":
    main()